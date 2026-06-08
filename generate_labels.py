import csv
import re
from collections import defaultdict
from docx import Document
from docx.shared import Pt
import openpyxl

# ---- helpers ----

def strip_letter_suffix(cid):
    m = re.match(r'^(\d+-\d+)[a-z]$', cid)
    return m.group(1) if m else cid

YEAR_PREFIX_MAP = {'01': 'Jan'}

def lookup_cid(base, coll_data):
    """Return (record, canonical_id) or (None, base). Handles 23-082 -> 23-82 and 01-XX -> Jan-XX."""
    if base in coll_data:
        return coll_data[base], base
    m = re.match(r'^(\d+)-0+(\d{2,})$', base)
    if m:
        normalized = f'{m.group(1)}-{m.group(2)}'
        if normalized in coll_data:
            return coll_data[normalized], normalized
    m = re.match(r'^(\d+)-(.+)$', base)
    if m:
        alias_prefix = YEAR_PREFIX_MAP.get(m.group(1))
        if alias_prefix:
            aliased = f'{alias_prefix}-{m.group(2)}'
            if aliased in coll_data:
                return coll_data[aliased], aliased
    return None, base

def id_sort_key(base):
    m = re.match(r'^(\w+)-(\d+)$', base)
    if m:
        prefix = m.group(1)
        try:
            return (int(prefix), int(m.group(2)))
        except ValueError:
            return (9999, int(m.group(2)))
    return (9999, 9999)

def format_date(raw):
    raw = raw.strip()
    m = re.search(r'(\d+)([ivxlcdm]+)(\d{4})', raw.lower())
    if m:
        day = str(int(m.group(1)))
        return f'{day}.{m.group(2)}.{m.group(3)}'
    return raw

# State and province abbreviation → full name
SUBDIVISION_NAMES = {
    # US states
    'AL': 'Alabama', 'AK': 'Alaska', 'AZ': 'Arizona', 'AR': 'Arkansas',
    'CA': 'California', 'CO': 'Colorado', 'CT': 'Connecticut', 'DE': 'Delaware',
    'FL': 'Florida', 'GA': 'Georgia', 'HI': 'Hawaii', 'ID': 'Idaho',
    'IL': 'Illinois', 'IN': 'Indiana', 'IA': 'Iowa', 'KS': 'Kansas',
    'KY': 'Kentucky', 'LA': 'Louisiana', 'ME': 'Maine', 'MD': 'Maryland',
    'MA': 'Massachusetts', 'MI': 'Michigan', 'MN': 'Minnesota', 'MS': 'Mississippi',
    'MO': 'Missouri', 'MT': 'Montana', 'NE': 'Nebraska', 'NV': 'Nevada',
    'NH': 'New Hampshire', 'NJ': 'New Jersey', 'NM': 'New Mexico', 'NY': 'New York',
    'NC': 'North Carolina', 'ND': 'North Dakota', 'OH': 'Ohio', 'OK': 'Oklahoma',
    'OR': 'Oregon', 'PA': 'Pennsylvania', 'RI': 'Rhode Island', 'SC': 'South Carolina',
    'SD': 'South Dakota', 'TN': 'Tennessee', 'TX': 'Texas', 'UT': 'Utah',
    'VT': 'Vermont', 'VA': 'Virginia', 'WA': 'Washington', 'WV': 'West Virginia',
    'WI': 'Wisconsin', 'WY': 'Wyoming',
    # Canadian provinces and territories
    'AB': 'Alberta', 'BC': 'British Columbia', 'MB': 'Manitoba',
    'NB': 'New Brunswick', 'NL': 'Newfoundland and Labrador',
    'NS': 'Nova Scotia', 'NT': 'Northwest Territories', 'NU': 'Nunavut',
    'ON': 'Ontario', 'PE': 'Prince Edward Island', 'QC': 'Quebec',
    'SK': 'Saskatchewan', 'YT': 'Yukon',
}

COUNTRY_NAMES = {
    'USA':           'United States of America',
    'CA':            'Canada',
    'United States': 'United States of America',
}

GEN_ROMAN = {'1': 'I', '2': 'II', '3': 'III', '4': 'IV', '5': 'V'}

# Gen label encodes morph only for gen 5 (genVapt / genValat); all other gens ignore morph
def gen_label(gen, morph):
    if gen == '5':
        if morph == 'aptera':
            return 'genVapt'
        if morph == 'alate':
            return 'genValat'
        return 'genV'
    return 'gen' + GEN_ROMAN.get(gen, gen)

GEN_LABEL_ORDER = ['genI', 'genII', 'genIV', 'genV', 'genVapt', 'genValat']

GEN_LABEL_DISPLAY = {
    'genI':    'gen. I fund.',
    'genII':   'gen. II al. vp.',
    'genIV':   'gen. IV apt. vp.',
    'genV':    'gen. V',
    'genVapt': 'gen. V apt. vp.',
    'genValat': 'gen. V al. sp.',
}

def parse_gps(s):
    s = (s or '').strip()
    if not s:
        return None
    parts = s.split(',')
    if len(parts) != 2:
        return None
    try:
        return (float(parts[0].strip()), float(parts[1].strip()))
    except ValueError:
        return None

def abbreviate_collectors(raw):
    """'Carol von Dohlen, Eric Maw' -> 'C. von Dohlen, E. Maw'"""
    if not raw:
        return raw
    names = [n.strip() for n in raw.split(',')]
    abbreviated = []
    for name in names:
        parts = name.split()
        if parts:
            parts[0] = parts[0][0] + '.'
        abbreviated.append(' '.join(parts))
    return ', '.join(abbreviated)

def abbreviate_collectors_and(raw):
    """Handle 'X and Y' or 'X, Y' collector strings → 'X. & Y.'"""
    if not raw:
        return raw
    parts = re.split(r'\s+and\s+|,\s*', raw)
    abbreviated = []
    for name in parts:
        name = name.strip()
        words = name.split()
        if words:
            words[0] = words[0][0] + '.'
        abbreviated.append(' '.join(words))
    return ' & '.join(abbreviated)

def county_location(location):
    """Strip site name after county, collapsing all sites in the same county to one key."""
    m = re.match(r'^(.*?\bCo\.),.*$', location)
    return m.group(1) if m else location

def normalize_location(raw):
    raw = raw.strip()
    m_dc = re.match(r'^USA: DC,\s*(.+)$', raw)
    if m_dc:
        return f'United States of America: Washington D.C., {m_dc.group(1)}'
    m = re.match(r'^([^:]+):\s*(.+)$', raw)
    if not m:
        return raw
    country_abbrev = m.group(1).strip()
    country = COUNTRY_NAMES.get(country_abbrev, country_abbrev)
    rest = m.group(2)
    parts = [p.strip() for p in rest.split(',')]
    subdivision = SUBDIVISION_NAMES.get(parts[0], parts[0])
    remaining = parts[1:]
    county_idx = None
    for i, p in enumerate(remaining):
        if re.search(r'\bCo\.', p):
            county_idx = i
            break
    if county_idx is not None:
        county = remaining[county_idx]
        site_parts = remaining[county_idx + 1:]
        site = ', '.join(site_parts)
        if site:
            return f'{country}: {subdivision}: {county}, {site}'
        return f'{country}: {subdivision}: {county}'
    site = ', '.join(remaining)
    if site:
        return f'{country}: {subdivision}, {site}'
    return f'{country}: {subdivision}'

NUM_WORDS = {
    1: 'One', 2: 'Two', 3: 'Three', 4: 'Four', 5: 'Five',
    6: 'Six', 7: 'Seven', 8: 'Eight', 9: 'Nine', 10: 'Ten',
}

def num_to_word(n):
    return NUM_WORDS.get(n, str(n))

def normalize_curation_note(raw):
    if not raw:
        return ''
    return raw[0].lower() + raw[1:]

# Gen+morph keyed lookups — morph values come from MorphoData.csv (aptera/alate),
# not from the Excel Morph column, so Gen I aptera = fundatrices (fp), not apt.
HIST_GEN_MORPH_HEADINGS = {
    ('1', 'aptera'):             'Generation I — Fundatrices',
    ('1', 'fundatrix'):          'Generation I — Fundatrices',
    ('2', 'alate'):              'Generation II — Alate Virginoparae',
    ('2', 'virginoparae alate'): 'Generation II — Alate Virginoparae',
    ('4', 'aptera'):             'Generation IV — Apterous Virginoparae',
    ('5', 'aptera'):             'Generation V — Apterous Virginoparae',
    ('5', 'alate'):              'Generation V — Alate Sexuparae',
}

HIST_GEN_MORPH_ABBREV = {
    ('1', 'aptera'):             'gen. I fund.',
    ('1', 'fundatrix'):          'gen. I fund.',
    ('2', 'alate'):              'gen. II al. vp.',
    ('2', 'virginoparae alate'): 'gen. II al. vp.',
    ('4', 'aptera'):             'gen. IV apt. vp.',
    ('5', 'aptera'):             'gen. V apt. vp.',
    ('5', 'alate'):              'gen. V al. sp.',
}

LECTOTYPE_PREFIXES = ('lectotype:', 'holotype:', 'paratype:', 'neotype:', 'paralectotype:', 'allotype:')

def hist_section_heading(gen, morph):
    key = (str(gen), str(morph).lower().strip())
    return HIST_GEN_MORPH_HEADINGS.get(key, f'Generation {gen} — {morph}')

# H and N are treated as one display group; L is separate
CLADE_GROUPS = [('L', ['L']), ('H+N', ['H', 'N'])]
ACTUAL_TO_DISPLAY_CLADE = {a: disp for disp, actuals in CLADE_GROUPS for a in actuals}

# ---- load historical data ----

hist_data = {}  # Collection_ID -> row dict
wb = openpyxl.load_workbook('../~data/HistoricCollData.xlsx')
ws = wb.active
hist_headers = [str(cell.value) for cell in next(ws.iter_rows(min_row=1, max_row=1))]
for row in ws.iter_rows(min_row=2, values_only=True):
    if row[0] is None:
        continue
    rec = dict(zip(hist_headers, row))
    cid = str(rec.get('Collection_ID', '')).strip()
    if cid:
        hist_data[cid] = rec

# ---- load data ----

# MorphoData.csv: count specimens per (base, gen, clade, morph)
slide_groups = defaultdict(int)

with open('../~data/MorphoData.csv', newline='', encoding='utf-8') as f:
    reader = csv.DictReader(f)
    for row in reader:
        base = strip_letter_suffix(row['Collection_ID'])
        key = (base, row['Gen'], row['Clade'], row['Morph'])
        slide_groups[key] += 1

# CollectionData.csv (local copy retains space-separated column names)
coll_data = {}
with open('CollectionData.csv', newline='', encoding='utf-8', errors='replace') as f:
    reader = csv.DictReader(f)
    for row in reader:
        coll_data[row['Collection ID']] = row

# ---- build entries, keyed by canonical ID to collapse alternate base formats ----
# (e.g. 23-082 and 23-82 both resolve to canonical 23-82 and must not become two entries)

canonical_gen_label_counts = defaultdict(lambda: defaultdict(int))
flagged_ids = set()
lectotype_entries = []

def _acc_str(raw):
    if isinstance(raw, (int, float)) and raw == raw:
        return str(int(raw))
    return str(raw).strip() if raw else ''

def _hist_id_from_rec(hist_rec, fallback_cid):
    """Return the canonical display ID for a historical slide."""
    raw = hist_rec.get('Historic ID')
    if raw is None or str(raw).strip() in ('', 'None', 'nan'):
        return str(hist_rec.get('Slide_ID', fallback_cid)).strip()
    if isinstance(raw, (int, float)) and raw == int(raw):
        return str(int(raw))
    return str(raw).strip()

def _hist_id_sort_key(hid):
    try:
        return (0, int(hid), '')
    except (ValueError, TypeError):
        return (1, 0, str(hid))

# hist_slide_specimens: keyed by Historic_ID, accumulates specimen count and metadata
hist_slide_specimens = defaultdict(lambda: {
    'count': 0, 'gen': '', 'morph': '', 'location': '', 'date': '',
    'host': '', 'collector': '', 'institution': '', 'accessions': set(), 'curation_notes': set(),
})

for (base, gen, clade, morph), count in slide_groups.items():
    rec, canonical_id = lookup_cid(base, coll_data)
    if rec is None:
        hist_rec = hist_data.get(base)
        if hist_rec is not None:
            # Type specimens: pull out separately, do not include in regular slide counts
            if any(str(base).lower().startswith(p) for p in LECTOTYPE_PREFIXES):
                lectotype_entries.append({
                    'type_status': str(base).split(':')[0].strip().capitalize(),
                    'gen':         str(gen),
                    'morph':       str(morph).lower().strip(),
                    'location':    normalize_location(str(hist_rec.get('Location', ''))),
                    'date':        format_date(str(hist_rec.get('Date', '?'))),
                    'host':        str(hist_rec.get('Host', '')),
                    'collector':   abbreviate_collectors_and(str(hist_rec.get('Collector', ''))),
                    'institution': str(hist_rec.get('Institution', '')),
                    'accession':   _acc_str(hist_rec.get('Collection Accession ID')),
                    'curation_note': normalize_curation_note(str(hist_rec.get('Special', '') or '')),
                })
                continue

            hist_id = _hist_id_from_rec(hist_rec, base)
            acc = _acc_str(hist_rec.get('Collection Accession ID'))
            note = normalize_curation_note(str(hist_rec.get('Special', '') or ''))

            slide = hist_slide_specimens[hist_id]
            slide['count'] += count
            slide['gen'] = str(gen)
            slide['morph'] = str(morph).lower().strip()
            slide['location'] = normalize_location(str(hist_rec.get('Location', '')))
            slide['date'] = format_date(str(hist_rec.get('Date', '?')))
            slide['host'] = str(hist_rec.get('Host', ''))
            slide['collector'] = abbreviate_collectors_and(str(hist_rec.get('Collector', '')))
            slide['institution'] = str(hist_rec.get('Institution', ''))
            if acc:
                slide['accessions'].add(acc)
            if note:
                slide['curation_notes'].add(note)
        else:
            flagged_ids.add(base)
        continue
    display_clade = ACTUAL_TO_DISPLAY_CLADE.get(clade, clade)
    canonical_gen_label_counts[(canonical_id, display_clade)][gen_label(gen, morph)] += count

entries = []
for (canonical_id, display_clade), glc in canonical_gen_label_counts.items():
    rec, _ = lookup_cid(canonical_id, coll_data)
    if rec is None:
        continue
    entries.append({
        'base_id': canonical_id,
        'gen_label_counts': dict(glc),
        'clade': display_clade,
        'location': normalize_location(rec['Location']),
        'date': format_date(rec['Date']),
        'host': rec['Host'],
        'collector': abbreviate_collectors(rec.get('Collector', '')),
        'gps': rec.get('Decimal Degrees GPS Coordinates', ''),
        'genbank': rec.get('GenBank COI', ''),
    })

# ---- group historical slides by collection event ----

hist_event_slides = defaultdict(list)
for hist_id, slide in hist_slide_specimens.items():
    event_key = (slide['gen'], slide['morph'], slide['location'], slide['date'], slide['host'], slide['collector'])
    hist_event_slides[event_key].append((hist_id, slide))

hist_events = []
for (gen, morph, location, date, host, collector), slides_list in hist_event_slides.items():
    total_count = sum(s['count'] for _, s in slides_list)
    institutions = sorted({s['institution'] for _, s in slides_list if s['institution']})
    accessions = set()
    curation_notes = set()
    for _, s in slides_list:
        accessions |= s['accessions']
        curation_notes |= s['curation_notes']
    sorted_ids = [hid for hid, _ in sorted(slides_list, key=lambda x: _hist_id_sort_key(x[0]))]
    hist_events.append({
        'gen': gen, 'morph': morph,
        'location': location, 'date': date, 'host': host, 'collector': collector,
        'total_count': total_count,
        'hist_ids': sorted_ids,
        'institutions': institutions,
        'accessions': accessions,
        'curation_notes': curation_notes,
    })

hist_events.sort(key=lambda e: (int(e['gen']) if e['gen'].isdigit() else 99, e['location']))

# ---- merge entries by (county_location, display_clade) ----

merge_buckets = defaultdict(list)
for e in entries:
    key = (county_location(e['location']), e['clade'])
    merge_buckets[key].append(e)

labels = []
for (county_loc, clade), group in merge_buckets.items():
    total_slides = sum(count for e in group for count in e['gen_label_counts'].values())

    # Sub-group by (date, host); sort entries within each sub-group by base_id
    sg_map = defaultdict(list)
    for e in group:
        sg_map[(e['date'], e['host'])].append(e)
    for k in sg_map:
        sg_map[k].sort(key=lambda e: id_sort_key(e['base_id']))

    # Build ordered sub-group list and collect discrepancies
    all_discrepancies = []
    sub_groups = []
    for (date, host) in sorted(sg_map.keys()):
        sg_entries = sg_map[(date, host)]
        collector = next((e['collector'] for e in sg_entries if e['collector']), '')

        sg_disc = []
        collector_vals = {e['base_id']: e['collector'] for e in sg_entries if e['collector']}
        if len(set(collector_vals.values())) > 1:
            detail = ', '.join(f'{b}: {v}' for b, v in sorted(collector_vals.items()))
            sg_disc.append(f'collector differs — {detail}')
        genbank_vals = {e['base_id']: e['genbank'] for e in sg_entries if e['genbank']}
        if len(set(genbank_vals.values())) > 1:
            detail = ', '.join(f'{b}: {v}' for b, v in sorted(genbank_vals.items()))
            sg_disc.append(f'GenBank differs — {detail}')
        all_discrepancies.extend(sg_disc)

        sub_groups.append({
            'date': date,
            'host': host,
            'collector': collector,
            'entries': sg_entries,
            'discrepancies': sg_disc,
        })

    labels.append({
        'location': county_loc,
        'clade': clade,
        'total_slides': total_slides,
        'sub_groups': sub_groups,
        'discrepancies': all_discrepancies,
    })

# ---- write document ----

def label_min_key(label):
    all_ids = [e['base_id'] for sg in label['sub_groups'] for e in sg['entries']]
    return min(id_sort_key(b) for b in all_ids)

def extract_country_state(location):
    """Return (country, state) from a normalized location string."""
    parts = location.split(': ', 2)
    if len(parts) < 2:
        return (location, '')
    country = parts[0]
    state = parts[1].split(',')[0].strip()
    return (country, state)

def format_gen_breakdown(gen_label_counts):
    def sort_key(k):
        try:
            return GEN_LABEL_ORDER.index(k)
        except ValueError:
            return 999
    sorted_labels = sorted(gen_label_counts.keys(), key=sort_key)
    parts = [f'{gen_label_counts[k]} {GEN_LABEL_DISPLAY.get(k, k)}' for k in sorted_labels]
    return '(' + ', '.join(parts) + ')'

def write_label_para(doc, lbl):
    p = doc.add_paragraph()
    p.add_run(f'{lbl["location"]}, ')

    prev_date = None
    prev_collector = None
    n_sgs = len(lbl['sub_groups'])

    for sg_i, sg in enumerate(lbl['sub_groups']):
        is_last_sg = (sg_i == n_sgs - 1)

        p.add_run('ex ')
        p.add_run(sg['host']).italic = True

        if sg['date'] != prev_date:
            p.add_run(f', {sg["date"]}')

        if sg['collector'] and sg['collector'] != prev_collector:
            p.add_run(f', {sg["collector"]}')

        p.add_run(', Coll. ID: ')

        n_entries = len(sg['entries'])
        for e_i, entry in enumerate(sg['entries']):
            is_last_entry = (e_i == n_entries - 1)
            breakdown = format_gen_breakdown(entry['gen_label_counts'])
            if is_last_entry and is_last_sg:
                p.add_run(f'{entry["base_id"]} {breakdown}.')
            elif is_last_entry:
                p.add_run(f'{entry["base_id"]} {breakdown}; ')
            else:
                p.add_run(f'{entry["base_id"]} {breakdown}, ')

        prev_date = sg['date']
        prev_collector = sg['collector']

def write_hist_label_para(doc, e):
    p = doc.add_paragraph()
    morph_label = HIST_GEN_MORPH_ABBREV.get((str(e['gen']), e['morph']), e['morph'])
    ids_str = ', '.join(str(hid) for hid in e['hist_ids'])

    p.add_run(f'{e["location"]}, ')
    p.add_run(e['host']).italic = True
    p.add_run(f', {e["date"]}, {e["collector"]}')
    p.add_run(f' ({e["total_count"]} {morph_label}; {ids_str})')

    curation_str = f' ({", ".join(sorted(e["curation_notes"]))})' if e['curation_notes'] else ''
    inst_str = ', '.join(e['institutions'])
    acc_str = ', '.join(sorted(e['accessions']))
    if acc_str:
        p.add_run(f'{curation_str} [{inst_str} acc. {acc_str}].')
    else:
        p.add_run(f'{curation_str} [{inst_str}].')

def write_lectotype_para(doc, e):
    p = doc.add_paragraph()
    morph_label = HIST_GEN_MORPH_ABBREV.get((str(e['gen']), e['morph']), e['morph'])
    run = p.add_run(e['type_status'])
    run.bold = True
    p.add_run(f' ♀ ({morph_label}): {e["location"]}, ')
    p.add_run(e['host']).italic = True
    p.add_run(f', {e["date"]}, {e["collector"]}')
    if e['curation_note']:
        p.add_run(f' ({e["curation_note"]})')
    if e['accession']:
        p.add_run(f' [{e["institution"]} acc. {e["accession"]}].')
    else:
        p.add_run(f' [{e["institution"]}].')

doc = Document()

label_index = defaultdict(list)
for lbl in labels:
    label_index[lbl['clade']].append(lbl)

def sort_labels(lbls):
    """Sort by: country → state (both by total slides desc), then label total slides desc, then lowest ID."""
    state_totals = defaultdict(int)
    country_totals = defaultdict(int)
    for lbl in lbls:
        cs = extract_country_state(lbl['location'])
        state_totals[cs] += lbl['total_slides']
        country_totals[cs[0]] += lbl['total_slides']
    lbls.sort(key=lambda lbl: (
        -country_totals[extract_country_state(lbl['location'])[0]],
        extract_country_state(lbl['location'])[0],
        -state_totals[extract_country_state(lbl['location'])],
        extract_country_state(lbl['location'])[1],
        -lbl['total_slides'],
        label_min_key(lbl),
    ))

first_section = True

for display_clade, _ in CLADE_GROUPS:
    combined = list(label_index.get(display_clade, []))
    if not combined:
        continue
    sort_labels(combined)

    if not first_section:
        doc.add_paragraph()

    doc.add_heading(f'Clade {display_clade}', level=1)
    first_section = False

    for lbl in combined:
        write_label_para(doc, lbl)

# ---- type material section ----

if lectotype_entries:
    doc.add_paragraph()
    doc.add_heading('Type Material', level=1)
    # Group by type status
    type_status_order = ['Holotype', 'Lectotype', 'Paratype', 'Paralectotype', 'Allotype', 'Neotype']
    by_status = defaultdict(list)
    for e in lectotype_entries:
        by_status[e['type_status']].append(e)
    for status in type_status_order:
        if status in by_status:
            doc.add_heading(status, level=2)
            for e in by_status[status]:
                write_lectotype_para(doc, e)
    for status, entries_list in by_status.items():
        if status not in type_status_order:
            doc.add_heading(status, level=2)
            for e in entries_list:
                write_lectotype_para(doc, e)

# ---- historical specimens section ----

if hist_events:
    doc.add_paragraph()
    doc.add_heading('Historical Morphometric Specimens', level=1)

    current_heading = None
    for e in hist_events:
        heading = hist_section_heading(e['gen'], e['morph'])
        if heading != current_heading:
            doc.add_heading(heading, level=2)
            current_heading = heading
        write_hist_label_para(doc, e)

# ---- flagged section ----

flagged_labels = [(lbl['location'], lbl['clade'], lbl['discrepancies'])
                  for lbl in labels if lbl['discrepancies']]

if flagged_labels:
    doc.add_paragraph()
    doc.add_heading('Flags — Data Discrepancies', level=1)
    for location, clade, discrepancies in flagged_labels:
        for disc in discrepancies:
            doc.add_paragraph(f'{location}, Clade {clade}: {disc}')

doc.add_paragraph()
doc.add_heading('Flagged — No Match in CollectionData.csv', level=1)
for fid in sorted(flagged_ids):
    doc.add_paragraph(fid)

doc.save('labels_output.docx')
print(f'Done. {len(labels)} modern labels written, {len(hist_events)} historical events, {len(lectotype_entries)} type specimens, {len(flagged_ids)} flagged IDs.')
